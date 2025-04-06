from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_resume():
    doc = Document()
    
    # Name
    name = doc.add_paragraph()
    name_run = name.add_run('JOHN SMITH')
    name_run.bold = True
    name_run.font.size = Pt(16)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact Info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run('Software Engineer\n')
    contact.add_run('Email: john.smith@email.com\n')
    contact.add_run('Phone: +1 (555) 123-4567')
    
    doc.add_paragraph()  # Spacing
    
    # Professional Summary
    doc.add_heading('PROFESSIONAL SUMMARY', level=1)
    doc.add_paragraph('Experienced Software Engineer with 6 years of expertise in developing scalable applications using Python and cloud technologies. Proven track record in leading technical projects and implementing machine learning solutions.')
    
    # Work Experience
    doc.add_heading('WORK EXPERIENCE', level=1)
    
    # Job 1
    p = doc.add_paragraph()
    p.add_run('Senior Software Developer').bold = True
    p.add_run(' | CloudTech Solutions | January 2022 - Present\n')
    doc.add_paragraph(
        '• Led development of microservices using Python and Docker\n'
        '• Implemented machine learning models for predictive analytics\n'
        '• Mentored junior developers and conducted code reviews\n'
        '• Managed AWS infrastructure for multiple projects'
    )
    
    # Job 2
    p = doc.add_paragraph()
    p.add_run('Software Engineer').bold = True
    p.add_run(' | DataInnovate | March 2019 - December 2021\n')
    doc.add_paragraph(
        '• Developed backend services using Python and SQL\n'
        '• Implemented agile methodologies and CI/CD pipelines\n'
        '• Created RESTful APIs using FastAPI and Django\n'
        '• Worked on AI-powered data analysis tools'
    )
    
    # Job 3
    p = doc.add_paragraph()
    p.add_run('Junior Software Developer').bold = True
    p.add_run(' | TechStart Inc | June 2017 - February 2019\n')
    doc.add_paragraph(
        '• Built web applications using JavaScript and Python\n'
        '• Maintained and optimized SQL databases\n'
        '• Collaborated in an agile development environment'
    )
    
    # Education
    doc.add_heading('EDUCATION', level=1)
    p = doc.add_paragraph()
    p.add_run('Bachelor of Science in Computer Science\n').bold = True
    p.add_run('Stanford University, 2017')
    
    # Skills
    doc.add_heading('SKILLS', level=1)
    doc.add_paragraph(
        '• Programming: Python, JavaScript, SQL\n'
        '• Technologies: AWS, Docker, Kubernetes\n'
        '• Frameworks: FastAPI, Django, React\n'
        '• Methodologies: Agile, Scrum\n'
        '• Other: Machine Learning, AI, Problem Solving'
    )
    
    # Certifications
    doc.add_heading('CERTIFICATIONS', level=1)
    doc.add_paragraph(
        '• AWS Certified Developer\n'
        '• Professional Scrum Master I'
    )
    
    # Save the document
    doc.save('data/sample_resume.docx')

if __name__ == '__main__':
    create_resume()
